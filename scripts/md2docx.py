#!/usr/bin/env python3
"""Convert project-proposal.md to Word (.docx) with proper Chinese formatting."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOCS_DIR = Path(__file__).resolve().parent.parent / "docs"
INPUT_MD = DOCS_DIR / "project-proposal.md"
OUTPUT_DOCX = DOCS_DIR / "project-proposal.docx"


def setup_styles(doc):
    """Configure document styles for Chinese content."""
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35

    for level in range(1, 5):
        sname = f"Heading {level}"
        if sname in doc.styles:
            hs = doc.styles[sname]
            hs.font.name = "微软雅黑"
            hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            sizes = {1: 22, 2: 16, 3: 14, 4: 12}
            hs.font.size = Pt(sizes[level])
            hs.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
            hs.paragraph_format.space_after = Pt(8 if level <= 2 else 6)


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Light blue background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "D6E4F0",
        })
        shading.append(shading_elm)

    # Data rows
    for r, row_data in enumerate(data_rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(10)
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()  # spacing after table


def parse_table_block(lines):
    """Parse markdown table lines into header + data rows."""
    # Filter out separator lines (|---|---|)
    content_lines = [l for l in lines if not re.match(r"^\|[\s\-:|]+\|$", l)]
    rows = []
    for line in content_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 1:
        return rows[0], rows[1:]
    return [], []


def add_code_block(doc, code_lines):
    """Add a code block as a formatted paragraph."""
    text = "\n".join(code_lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Light gray background via shading
    shading = run.element.get_or_add_rPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F5F5F5",
    })
    shading.append(shading_elm)


def process_inline(paragraph, text):
    """Process inline markdown (bold, code, links) and add runs."""
    # Pattern: **bold**, `code`, [text](url), plain text
    pattern = r"(\*\*(.+?)\*\*|`(.+?)`|\[(.+?)\]\(.+?\)|([^*`\[]+))"
    for match in re.finditer(pattern, text):
        full = match.group(0)
        if match.group(2):  # bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # code
            run = paragraph.add_run(match.group(3))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif match.group(4):  # link text
            run = paragraph.add_run(match.group(4))
            run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
            run.underline = True
        else:  # plain
            run = paragraph.add_run(full)
        run.font.name = "微软雅黑"
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def convert(input_path, output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    setup_styles(doc)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush table if pending
                if in_table:
                    header, data = parse_table_block(table_lines)
                    if header:
                        add_table(doc, header, data)
                    table_lines = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if re.match(r"^\|.+\|$", line.strip()):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            header, data = parse_table_block(table_lines)
            if header:
                add_table(doc, header, data)
            table_lines = []
            in_table = False

        # Skip horizontal rules
        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Title (H1) - centered
            if level == 1:
                p = doc.add_heading(text, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level=level)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            text = re.sub(r"^>\s*", "", line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.name = "微软雅黑"
            run.font.size = Pt(11)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^(\s*)[-*·]\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            level = min(indent // 2, 2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1 + level * 0.8)
            process_inline(p, text)
            # Clear the auto-generated text
            if p.runs:
                pass  # runs already added by process_inline
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        process_inline(p, line.strip())
        i += 1

    # Flush remaining table
    if in_table and table_lines:
        header, data = parse_table_block(table_lines)
        if header:
            add_table(doc, header, data)

    doc.save(str(output_path))
    print(f"Generated: {output_path}")


if __name__ == "__main__":
    convert(INPUT_MD, OUTPUT_DOCX)
